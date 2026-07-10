from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


@dataclass
class LoadedPage:
    text: str
    page: int | None = None


@dataclass
class LoadedDocument:
    title: str
    source: str
    pages: list[LoadedPage]


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_file(path: Path, title: str | None = None, source: str | None = None) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".pdf":
        pages = _load_pdf(path)
    elif suffix == ".docx":
        pages = _load_docx(path)
    elif suffix in {".html", ".htm"}:
        pages = [LoadedPage(_html_to_text(path.read_text(encoding="utf-8", errors="ignore")))]
    else:
        pages = [LoadedPage(clean_text(path.read_text(encoding="utf-8", errors="ignore")))]

    return LoadedDocument(title=title or path.stem, source=source or str(path), pages=pages)


def load_url(url: str, title: str | None = None, timeout: int = 45) -> LoadedDocument:
    response = requests.get(url, timeout=timeout, headers={"User-Agent": "finance-rag-assistant/1.0"})
    response.raise_for_status()
    content_type = response.headers.get("content-type", "").lower()
    name = title or url.rstrip("/").split("/")[-1] or url

    if "pdf" in content_type or url.lower().endswith(".pdf"):
        reader = PdfReader(io.BytesIO(response.content))
        pages = [
            LoadedPage(clean_text(page.extract_text() or ""), page=i + 1)
            for i, page in enumerate(reader.pages)
        ]
    elif "officedocument.wordprocessingml" in content_type or url.lower().endswith(".docx"):
        doc = Document(io.BytesIO(response.content))
        pages = [LoadedPage(clean_text("\n".join(p.text for p in doc.paragraphs)))]
    else:
        pages = [LoadedPage(_html_to_text(response.content))]

    return LoadedDocument(title=name, source=url, pages=[p for p in pages if p.text])


def chunk_document(doc: LoadedDocument, chunk_size: int = 900, overlap: int = 160) -> Iterable[dict]:
    for page in doc.pages:
        text = clean_text(page.text)
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            boundary = max(text.rfind("\n", start, end), text.rfind("。", start, end), text.rfind(".", start, end))
            if boundary > start + chunk_size * 0.55:
                end = boundary + 1
            chunk_text = clean_text(text[start:end])
            if len(chunk_text) >= 80:
                yield {"text": chunk_text, "page": page.page}
            if end >= len(text):
                break
            start = max(0, end - overlap)


def _load_pdf(path: Path) -> list[LoadedPage]:
    reader = PdfReader(str(path))
    return [
        LoadedPage(clean_text(page.extract_text() or ""), page=i + 1)
        for i, page in enumerate(reader.pages)
    ]


def _load_docx(path: Path) -> list[LoadedPage]:
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return [LoadedPage(clean_text(text))]


def _html_to_text(html: str | bytes) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return clean_text(soup.get_text("\n"))
